"""
AI Interview Feedback MVP - Resume Service

This module handles resume processing and analysis:
- Text extraction from PDF and DOCX files
- Skill extraction from resume text
- Resume-JD comparison and skill matching

Author: Member 2 (ML Engine)

Key Functions:
    - extract_text_from_resume(path): Extract text from PDF/DOCX
    - extract_skills(text): Extract skill keywords from text
    - compare_resume_with_jd(resume_text, jd_text): Compare skills
"""

import os
import re
from typing import Dict, List, Set
from app.config import COMMON_SKILLS


# ===========================================
# Text Extraction
# ===========================================

def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text content from a resume file.
    
    Supports PDF and DOCX formats. Automatically detects
    file type based on extension.
    
    Args:
        file_path: Path to the resume file
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If file format is not supported
    
    Example:
        >>> text = extract_text_from_resume("resume.pdf")
        >>> print(text[:100])
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return _extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2.
    
    Args:
        file_path: Path to the PDF file
    
    Returns:
        str: Extracted text
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""


def _extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        str: Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""


# ===========================================
# Skill Extraction
# ===========================================

def extract_skills(text: str) -> Set[str]:
    """
    Extract skill keywords from text.
    
    Searches for common technical and soft skills
    defined in the configuration.
    
    Args:
        text: Text to extract skills from
    
    Returns:
        Set[str]: Set of found skill keywords
    
    Example:
        >>> skills = extract_skills("Experience with Python, SQL, and machine learning")
        >>> print(skills)  # {'python', 'sql', 'machine learning'}
    """
    if not text:
        return set()
    
    text_lower = text.lower()
    found_skills = set()
    
    for skill in COMMON_SKILLS:
        # Check if skill appears in text (word boundary aware for short skills)
        if len(skill) <= 3:
            # For short skills like "sql", "aws", use word boundary
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill)
        else:
            # For longer skills, simple contains check
            if skill in text_lower:
                found_skills.add(skill)
    
    return found_skills


def extract_skills_advanced(text: str) -> Dict[str, List[str]]:
    """
    Extract skills from text with categorization.
    
    Returns skills grouped by category:
    - programming: Programming languages
    - frameworks: Web/ML frameworks
    - databases: Database technologies
    - cloud: Cloud/DevOps tools
    - soft_skills: Soft skills
    
    Args:
        text: Text to extract skills from
    
    Returns:
        Dict[str, List[str]]: Skills grouped by category
    """
    skill_categories = {
        "programming": ["python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php"],
        "frameworks": ["react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi", "tensorflow", "pytorch"],
        "databases": ["sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch"],
        "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "ci/cd", "jenkins", "terraform"],
        "soft_skills": ["leadership", "communication", "teamwork", "problem solving", "agile", "scrum"]
    }
    
    text_lower = text.lower()
    result = {category: [] for category in skill_categories}
    
    for category, skills in skill_categories.items():
        for skill in skills:
            pattern = r'\b' + re.escape(skill.replace('.', r'\.')) + r'\b'
            if re.search(pattern, text_lower, re.IGNORECASE):
                result[category].append(skill)
    
    return result


# ===========================================
# Resume-JD Comparison
# ===========================================

def compare_resume_with_jd(resume_text: str, jd_text: str) -> Dict:
    """
    Compare a resume against a job description.
    
    Extracts skills from both documents and calculates:
    - Matched skills (in both)
    - Missing skills (in JD but not resume)
    - Match percentage
    - Semantic similarity
    
    Args:
        resume_text: Extracted text from resume
        jd_text: Job description text
    
    Returns:
        dict: Comparison results:
            {
                "skill_match_pct": float (0-100),
                "matched_skills": List[str],
                "missing_skills": List[str],
                "extra_skills": List[str],
                "similarity_score": float (0-1)
            }
    
    Example:
        >>> result = compare_resume_with_jd(resume_text, jd_text)
        >>> print(f"Match: {result['skill_match_pct']}%")
    """
    # Extract skills from both documents
    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(jd_text)
    
    # Calculate matches
    matched_skills = resume_skills & jd_skills
    missing_skills = jd_skills - resume_skills
    extra_skills = resume_skills - jd_skills
    
    # Calculate match percentage
    if len(jd_skills) == 0:
        # If no skills found in JD, use semantic similarity instead
        skill_match_pct = 50.0  # Neutral score
    else:
        skill_match_pct = (len(matched_skills) / len(jd_skills)) * 100
    
    # Get semantic similarity for additional context
    from app.services.ml_engine import semantic_similarity
    similarity_score = semantic_similarity(resume_text, jd_text)
    
    # Combine keyword match and semantic similarity
    # Weight: 60% keyword match, 40% semantic similarity
    combined_score = skill_match_pct * 0.6 + (similarity_score * 100) * 0.4
    
    return {
        "skill_match_pct": round(combined_score, 1),
        "matched_skills": sorted(list(matched_skills)),
        "missing_skills": sorted(list(missing_skills)),
        "extra_skills": sorted(list(extra_skills)),
        "similarity_score": round(similarity_score, 3),
        "jd_skills_count": len(jd_skills),
        "resume_skills_count": len(resume_skills)
    }


# ===========================================
# Additional Analysis Functions
# ===========================================

def analyze_resume_sections(text: str) -> Dict[str, bool]:
    """
    Check for presence of common resume sections.
    
    Args:
        text: Resume text
    
    Returns:
        Dict[str, bool]: Section presence indicators
    """
    sections = {
        "contact_info": False,
        "summary": False,
        "experience": False,
        "education": False,
        "skills": False,
        "projects": False,
        "certifications": False
    }
    
    text_lower = text.lower()
    
    # Check for each section
    section_keywords = {
        "contact_info": ["email", "phone", "linkedin", "github"],
        "summary": ["summary", "objective", "profile", "about"],
        "experience": ["experience", "employment", "work history", "professional"],
        "education": ["education", "university", "college", "degree", "bachelor", "master"],
        "skills": ["skills", "technologies", "technical", "competencies"],
        "projects": ["projects", "portfolio"],
        "certifications": ["certifications", "certificates", "certification"]
    }
    
    for section, keywords in section_keywords.items():
        for keyword in keywords:
            if keyword in text_lower:
                sections[section] = True
                break
    
    return sections


def get_resume_word_count(text: str) -> Dict[str, int]:
    """
    Get word count statistics for a resume.
    
    Args:
        text: Resume text
    
    Returns:
        Dict containing word count statistics
    """
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    
    return {
        "word_count": len(words),
        "sentence_count": len([s for s in sentences if s.strip()]),
        "avg_words_per_sentence": len(words) / max(1, len(sentences))
    }
